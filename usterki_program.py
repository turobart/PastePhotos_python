# -*- coding: utf-8 -*-

import wx
import os.path
import math
from PIL import Image
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK

import wx.lib.scrolledpanel as scrolled
import shutil
import sys


size_x=700+70+10
size_y=600
global photos_paths, items_list, list_1st_photo, list_2nd_photo, list_3rd_photo, list_4th_photo
global list_1st_txt, list_2nd_txt, list_3rd_txt, list_4th_txt, photo_count, photos, flaw_list, photo_directory, list_progress_count, prgrs_wndw, result_dir
global small_photos_dir_path, used_photos_dir, protocol_name, protocol_path, small_photos_list, protocol_saved, saved_photos, data_path
global work_dir
saved_photos=[]
protocol_saved=True
small_photos_list=[]
protocol_path=''
protocol_name=''
small_photos_dir_path=None
used_photos_dir=None
prgrs_wndw=None
list_progress_count=0
flaw_list=[]
photos=[]
photo_count=0
list_1st_photo=[]
list_2nd_photo=[]
list_3rd_photo=[]
list_1st_txt=[]
list_2nd_txt=[]
list_3rd_txt=[]
photos_paths=[]
items_list=[]

flaw_cell_size=(120,40)

dir_path=os.path.dirname(os.path.realpath(__file__))
photo_directory=dir_path

save_path=os.path.join(dir_path, 'zapisane')
work_dir=os.path.join(dir_path, 'roboczy')
result_dir=os.path.join(dir_path, 'wyniki')

if not os.path.exists(save_path):
    os.makedirs(save_path)
if not os.path.exists(work_dir):
    os.makedirs(work_dir)
if not os.path.exists(result_dir):
    os.makedirs(result_dir)

data_path = getattr(sys, '_MEIPASS', os.getcwd())
 
ID_ADD=1
ID_REMOVE=2
ID_2=3
ID_3=4
ID_OPT=5
ID_START=6
ID_SAVE=7
ID_OPEN=8
ID_NEW=9
ID_SAVE_AS=10
ID_D=11
ID_D_Y=12
ID_D_N=13
ID_PPP=14
BT_LFT_ID=501
BT_DEL_ID=502
BT_RHT_ID=503
BT_ADD_ID=504
ID_3x2 = 505

class main_window(wx.Frame):
    def __init__(self, parent, title):
        super(main_window, self).__init__(parent, title=title, size=(size_x, size_y))
          
        self.InitUI()    
        self.Centre()
        self.Show()    
        
    def InitUI(self):
        
        self.menubar = wx.MenuBar()
        self.mainMenu = wx.Menu()
        
        self.start_new = wx.MenuItem(self.mainMenu, ID_NEW, 'Nowa lista\tCtrl+N')
        self.menu_open_button = wx.MenuItem(self.mainMenu, ID_OPEN, 'Otwórz\tCtrl+Q')
        self.menu_save_button = wx.MenuItem(self.mainMenu, ID_SAVE, 'Zapisz\tCtrl+S')
        self.menu_save_as_button = wx.MenuItem(self.mainMenu, ID_SAVE_AS, 'Zapisz jako...')
        self.add_next = wx.MenuItem(self.mainMenu, ID_ADD, 'Dodaj usterkę\tCtrl+D')
        self.remove_last = wx.MenuItem(self.mainMenu, ID_REMOVE, 'Usuń ostatnią usterkę\tCtrl+U')
        self.menu_start_button = wx.MenuItem(self.mainMenu, ID_START, 'Rozpocznij\tCtrl+R')
        
        self.Bind(wx.EVT_MENU, self.start_project, id=ID_NEW)
        self.Bind(wx.EVT_MENU, self.add_new, id=ID_ADD)
        self.Bind(wx.EVT_MENU, self.remove_old, id=ID_REMOVE)
        self.Bind(wx.EVT_MENU, self.start_clicked, id=ID_START)
        self.Bind(wx.EVT_MENU, self.save_all, id=ID_SAVE)
        self.Bind(wx.EVT_MENU, self.save_as, id=ID_SAVE_AS)
        self.Bind(wx.EVT_MENU, self.open_old, id=ID_OPEN)
        
        self.menu_save_button.Enable(False)
        self.menu_save_as_button.Enable(False)
        self.add_next.Enable(False)
        self.remove_last.Enable(False)
        self.menu_start_button.Enable(False)
        
        self.optionsMenu = wx.Menu(ID_OPT)
        self.photo_per_page = wx.Menu(ID_PPP)
        self.two_photo=wx.MenuItem(self.photo_per_page, ID_2, '2', kind=wx.ITEM_RADIO)
        self.three_photo=wx.MenuItem(self.photo_per_page, ID_3, '3', kind=wx.ITEM_RADIO)
        self.six_photo=wx.MenuItem(self.photo_per_page, ID_3x2, '3x2', kind=wx.ITEM_RADIO)
        self.photo_per_page.Append(self.two_photo)
        self.photo_per_page.Append(self.three_photo)
        self.photo_per_page.Append(self.six_photo)
        self.six_photo.Check(check=True)
        
        self.description_menu = wx.Menu(ID_D)
        self.description_yes=wx.MenuItem(self.description_menu, ID_D_Y, 'Tak', kind=wx.ITEM_RADIO)
        self.description_no=wx.MenuItem(self.description_menu, ID_D_N, 'Nie', kind=wx.ITEM_RADIO)
        self.description_menu.Append(self.description_yes)
        self.description_menu.Append(self.description_no)
        self.description_yes.Check(check=True)
        
        self.optionsMenu.Append(ID_PPP, 'Ilość zdjęć na stonie', self.photo_per_page)
        self.optionsMenu.Append(ID_D, 'Podpis pod zdjęciem: tak', self.description_menu)
        self.optionsMenu.Enable(ID_D, False)
        
        self.Bind(wx.EVT_MENU, self.description_yes_no, id=ID_D_Y)
        self.Bind(wx.EVT_MENU, self.description_yes_no, id=ID_D_N)
        
        self.mainMenu.Append(self.start_new)
        self.mainMenu.Append(self.menu_open_button)
        self.mainMenu.Append(self.menu_save_button)
        self.mainMenu.Append(self.menu_save_as_button)
        self.mainMenu.AppendSeparator()
        self.mainMenu.Append(self.add_next)
        self.mainMenu.Append(self.remove_last)
        self.mainMenu.AppendSeparator()
        self.mainMenu.Append(self.menu_start_button)
        
        self.menubar.Append(self.mainMenu, '&Plik')
        self.menubar.Append(self.optionsMenu, '&Opcje')
        self.SetMenuBar(self.menubar)
        
        dummy_panel=wx.Panel(self)
#         project_panel=wx.Panel(dummy_panel, size=(size_x-300, 25), style=wx.SIMPLE_BORDER)
        
#         self.project_name=wx.StaticText(project_panel, label='Projekt: ', style=wx.ALIGN_CENTER, size=(size_x-300,20))
        
        info_panel=wx.Panel(dummy_panel, size=(size_x-70, 35), style=wx.SIMPLE_BORDER)
        self.info_sizer=wx.FlexGridSizer(1,8,10,10)
        
        self.flaw_nr=wx.StaticText(info_panel, label='Nr', style=wx.ALIGN_CENTER, size=(20,30))
        self.flaw_text=wx.StaticText(info_panel, label='Usterka', style=wx.ALIGN_CENTER, size=(flaw_cell_size[0],30))
        self.confirmation_text=wx.StaticText(info_panel, label='', style=wx.ALIGN_CENTER, size=(70,35))
        
        self.flaw_level_text=wx.StaticText(info_panel, label='Stopień pilności', style=wx.ALIGN_CENTER, size=(70,30))
        self.flaw_order_text=wx.StaticText(info_panel, label='Zmiana kolejności', style=wx.ALIGN_CENTER, size=(150,30))
        self.flaw_important_text=wx.StaticText(info_panel, label='Wyróżnienie', style=wx.ALIGN_CENTER, size=(70,30))
        self.check_photos_text=wx.StaticText(info_panel, label='Sprawdź zdjęcia', style=wx.ALIGN_CENTER, size=(70,30))
        self.disable_enable_text=wx.StaticText(info_panel, label='Bez zdjęć', style=wx.ALIGN_CENTER, size=(70,30))
        
        self.info_sizer.Add(self.flaw_nr)
        self.info_sizer.Add(self.flaw_text)
        self.info_sizer.Add(self.confirmation_text)
        self.info_sizer.Add(self.flaw_level_text)
        self.info_sizer.Add(self.flaw_order_text)
        self.info_sizer.Add(self.flaw_important_text)
        self.info_sizer.Add(self.check_photos_text)
        self.info_sizer.Add(self.disable_enable_text)
        
        
        info_panel.SetSizer(self.info_sizer)
        info_panel.Layout()
        
        self.main_panel = wx.Panel(dummy_panel, size=(size_x-50, size_y+200))
        self.main_panel.SetBackgroundColour('#4f5049')
        self.main_sizer=wx.BoxSizer(wx.VERTICAL)
        
        
        
        self.scroll_panel=scrolled.ScrolledPanel(self.main_panel, size=(size_x-50, size_y+200))
        self.scroll_panel.SetAutoLayout(1)
        self.scroll_panel.SetupScrolling(scroll_x=False)
        self.scroll_panel.SetBackgroundColour('#ededed')
        self.scroll_sizer=wx.BoxSizer(wx.VERTICAL)
        self.main_sizer.Add(self.scroll_panel)
        self.scroll_panel.SetSizer(self.scroll_sizer)
        
        dummy_sizer=wx.BoxSizer(wx.VERTICAL)
#         dummy_sizer.Add(project_panel, flag=wx.TOP|wx.ALIGN_CENTER, border=10)
        dummy_sizer.Add(info_panel, flag=wx.LEFT|wx.TOP, border=10)
        dummy_sizer.Add(self.main_panel, flag=wx.ALL, border=10)
        
        os.chdir(data_path)
        ico=wx.Icon('zdjecia.ico', wx.BITMAP_TYPE_ICO)
        self.SetIcon(ico)
        self.Bind(wx.EVT_CLOSE, self.CloseSelf)
        
        self.main_panel.SetSizer(self.main_sizer)

        dummy_panel.SetSizer(dummy_sizer)
        
        self.SetMinSize((size_x-15, size_y-200))
        self.SetMaxSize((size_x-15, size_y+200))
    
    def CloseSelf(self, event):
        self.save_quit()
        
        
    def save_quit(self):
        global save_path, protocol_name, protocol_path, items_list, photo_directory, protocol_saved, saved_photos, used_photos_dir
        if protocol_name=='' or protocol_saved or protocol_name[0]==' ':
            self.Destroy()
        else:            
            path=os.path.join(save_path, protocol_name+'.btp')
            save_quit_message='Czy na zapisać przed wyjściem?'
            save_quit_dlg=wx.MessageBox(save_quit_message, 'Potwierdzenie', wx.YES_NO|wx.CANCEL) 
            if save_quit_dlg==wx.YES:
                if os.path.exists(path):
                    overwrite_save_quit_message='Plik o podanej nazwie istnieje. Nadpisać?'
                    overwrite_save_quit_dlg=wx.MessageBox(overwrite_save_quit_message, 'Ostrzeżenie', wx.YES_NO|wx.CANCEL) 
                    if overwrite_save_quit_dlg==wx.NO:
                        all_photos=[]
                        for item in items_list:
                            if item.old_photos:
                                for old_photo in item.old_photos:
                                    all_photos.append(old_photo)
                        for photo in saved_photos:
                            if not os.path.exists(photo):
                                find_photo=os.path.join(photo_directory,os.path.basename(photo))
                                if os.path.exists(find_photo):
                                    shutil.move(find_photo, used_photos_dir)
                                elif photo_directory is not dir_path:
                                    find_photo=os.path.join(dir_path,os.path.basename(photo))
                                    if os.path.exists(find_photo):
                                        shutil.move(find_photo, used_photos_dir)
                        photos_to_move=[item1 for item1 in all_photos if item1 not in saved_photos]
                        for photo in photos_to_move:
                            shutil.move(photo,photo_directory)
#                             item.old_photos=[]
                        self.Destroy()
                        return
                    elif overwrite_save_quit_dlg==wx.CANCEL:
                        return
                    
                photo_per_page_item=self.photo_per_page.GetMenuItems()
        
                for item_nr in photo_per_page_item:
                    nr_bool=self.photo_per_page.IsChecked(item_nr.GetId())
                    if nr_bool == True:
                        photo_radio_id=item_nr.GetId()    
                    
                with open(path, "w") as myfile:
                    myfile.write(photo_directory+'\n')
                    myfile.write(str(photo_radio_id)+'\n')
                    myfile.write(str(self.description_yes.IsChecked())+'\n\n')
                    
                    for item in items_list:
                        myfile.write(item.text_panel.GetLineText(0)+'\t')
                    myfile.write('\n')
                    for item in items_list:
                        for element in item.old_photos:
                            element_dir=os.path.basename(os.path.dirname(element))
                            project_dir=os.path.basename(os.path.dirname(os.path.dirname(element)))
                            final_element_path=os.path.join(project_dir, element_dir, os.path.basename(element))
                            myfile.write(str(final_element_path)+',')
                        myfile.write('\t')
                    myfile.write('\n')
                    for item in items_list:
                        myfile.write(str(item.photo_choice.GetSelection())+'\t')
                    myfile.write('\n')
                    for item in items_list:
                        myfile.write(str(item.check_box.GetValue())+'\t')
                    myfile.write('\n')
                    for item in items_list:
                        myfile.write(str(item.no_foto_cb.GetValue())+'\t')
                shutil.copy(path,protocol_path)
                self.Destroy()
            elif save_quit_dlg==wx.CANCEL:
                return
            else:
                all_photos=[]
                for item in items_list:
                    if item.old_photos:
                        for old_photo in item.old_photos:
                            all_photos.append(old_photo)
                for photo in saved_photos:
                    if not os.path.exists(photo):
                        find_photo=os.path.join(photo_directory,os.path.basename(photo))
                        if os.path.exists(find_photo):
                            shutil.move(find_photo, used_photos_dir)
                        elif photo_directory is not dir_path:
                            find_photo=os.path.join(dir_path,os.path.basename(photo))
                            if os.path.exists(find_photo):
                                shutil.move(find_photo, used_photos_dir)
                photos_to_move=[item1 for item1 in all_photos if item1 not in saved_photos]
                for photo in photos_to_move:
                    shutil.move(photo,photo_directory)
                self.Destroy()
        
        
    def description_yes_no(self, event):
        to_find='nie'
        des_str=self.optionsMenu.GetLabel(ID_D).find(to_find, 20, 24)
        if des_str<0:
            self.optionsMenu.SetLabel(ID_D, 'Podpis pod zdjęciem: nie')
        else:
            self.optionsMenu.SetLabel(ID_D, 'Podpis pod zdjęciem: tak') 
         
    def start_clicked(self, event):
        global list_1st_photo, list_2nd_photo, list_3rd_photo, list_4th_photo
        global list_1st_txt, list_2nd_txt, list_3rd_txt, list_4th_txt, photos, photo_count, flaw_list, prgrs_wndw,list_progress_count
        global list_1st_no_photo_txt, list_2nd_no_photo_txt, list_3rd_no_photo_txt, list_4th_no_photo_txt
        start_message=''
        text_item_pos=[]
        photo_item_pos=[]
        flaw_item_pos=[]
        text_bool=False
        photo_bool=False
        flaw_bool=False
        for item in items_list:
            if item.text_panel.GetLineText(0)=='':
                text_item_pos.append(items_list.index(item)+1)
                text_bool=True
            if len(item.old_photos)==0 and not item.no_foto_cb.GetValue():
                photo_item_pos.append(items_list.index(item)+1)
                photo_bool=True
            if item.photo_choice.GetSelection()==-1:
                flaw_item_pos.append(items_list.index(item)+1)
                flaw_bool=True
        if text_bool or photo_bool or flaw_bool:
            if text_bool:start_message+=('W usterce: ' + ', '.join('{}'.format(k[1]) for k in enumerate(text_item_pos)) + ' brak opisu.\n')
            if photo_bool: start_message+=('W usterce: ' + ', '.join('{}'.format(k[1]) for k in enumerate(photo_item_pos)) + ' brak zdjęć.\n')
            if flaw_bool: start_message+=('W usterce: ' + ', '.join('{}'.format(k[1]) for k in enumerate(flaw_item_pos)) + ' brak stopnia pilności.')            
            starts_dlg=wx.MessageBox(start_message, 'Ostrzeżenie', wx.OK) 
            return
        
        list_progress_count=0
        photo_progress_max=0
        flaw_list=[]
        photo_count=0
        photos=[]
        list_1st_photo=[]
        list_1st_txt=[]
        list_2nd_photo=[]
        list_2nd_txt=[]
        list_3rd_photo=[]
        list_3rd_txt=[]
        list_4th_photo=[]
        list_4th_txt=[]
        list_1st_no_photo_txt=[]
        list_2nd_no_photo_txt=[]
        list_3rd_no_photo_txt=[]
        list_4th_no_photo_txt=[]
        for item in items_list:
            photo_progress_max+=len(item.old_photos)
            if item.photo_choice.GetSelection()==0:
                text=item.text_panel.GetLineText(0)
                if item.no_foto_cb.GetValue():
                    list_1st_no_photo_txt.append(text)
                else:
                    if item.photos_paths: list_1st_photo.append(item.photos_paths)
                    if text: list_1st_txt.append(text)
            elif item.photo_choice.GetSelection()==1:
                text=item.text_panel.GetLineText(0)
                if item.no_foto_cb.GetValue():
                    list_2nd_no_photo_txt.append(text)
                else:
                    if item.photos_paths: list_2nd_photo.append(item.photos_paths)
                    if text: list_2nd_txt.append(text)
            elif item.photo_choice.GetSelection()==2:
                text=item.text_panel.GetLineText(0)
                if item.no_foto_cb.GetValue():
                    list_3rd_no_photo_txt.append(text)
                else:
                    if item.photos_paths: list_3rd_photo.append(item.photos_paths)
                    if text: list_3rd_txt.append(text)
            elif item.photo_choice.GetSelection()==3:
                text=item.text_panel.GetLineText(0)
                if item.no_foto_cb.GetValue():
                    list_4th_no_photo_txt.append(text)
                else:
                    if item.photos_paths: list_4th_photo.append(item.photos_paths)
                    if text: list_4th_txt.append(text)
        
#         len(list_1st_photo)+len(list_2nd_photo)+len(list_3rd_photo)
        
        prgrs_wndw=progress_window(main_frame)
        prgrs_wndw.list_progress.SetRange(len(list_1st_txt)+len(list_2nd_txt)+len(list_3rd_txt)+len(list_4th_txt)+len(list_1st_no_photo_txt)+len(list_2nd_no_photo_txt)+len(list_3rd_no_photo_txt)+len(list_4th_no_photo_txt))
        prgrs_wndw.photo_progress.SetRange(photo_progress_max)
        
        
        self.confirmation_text.SetLabel('')
        
        flaws_list()
        photo_table()
        
        prgrs_wndw.dlg_info_text.SetLabel('Lista i tabela zapisane w\nfolderze "wyniki"')
            
    def add_new(self, event):
        global protocol_saved
        protocol_saved=False
        enable_save()
        enable_start()
        new=single(self.scroll_panel)
        self.scroll_sizer.Add(new, flag=wx.DOWN, border=5)
        items_list.append(new)
        self.main_panel.Layout()
        
        self.confirmation_text.SetLabel('')
        self.confirmation_text.SetBackgroundColour('#f0f0f0')
        
    def remove_old(self, event):
        global protocol_saved
        child_list=self.scroll_sizer.GetChildren()
        if child_list:
            remove_message='Czy na pewno usunąć ostatnią usterkę?'
            remove_dlg=wx.MessageBox(remove_message, 'Potwierdzenie', wx.YES_NO) 
            if remove_dlg==wx.YES:
                if items_list[-1].old_photos:
                    for old_photo in items_list[-1].old_photos:
                        shutil.move(old_photo,photo_directory)
                self.scroll_sizer.Hide(len(child_list)-1)
                self.scroll_sizer.Remove(len(child_list)-1)
                self.scroll_panel.Layout()
                if len(child_list)<11:
                    self.scroll_panel.Fit()
                del items_list[-1]
                self.main_panel.Layout()
                protocol_saved=False
        self.confirmation_text.SetLabel('')
        self.confirmation_text.SetBackgroundColour('#f0f0f0')

    def save_all(self,event):
        global protocol_name, protocol_path, protocol_saved, saved_photos, photo_directory
        path=os.path.join(save_path, protocol_name+'.btp')
        protocol_saved=True
        photo_per_page_item=self.photo_per_page.GetMenuItems()
        
        for item_nr in photo_per_page_item:
            nr_bool=self.photo_per_page.IsChecked(item_nr.GetId())
            if nr_bool == True:
                photo_radio_id=item_nr.GetId()
        
        with open(path, "w") as myfile:
            myfile.write(photo_directory+'\n')
            myfile.write(str(photo_radio_id)+'\n')
            myfile.write(str(self.description_yes.IsChecked())+'\n\n')
            
            for item in items_list:
                myfile.write(item.text_panel.GetLineText(0)+'\t')
            myfile.write('\n')
            for item in items_list:
                for element in item.old_photos:
                    element_dir=os.path.basename(os.path.dirname(element))
                    project_dir=os.path.basename(os.path.dirname(os.path.dirname(element)))
                    final_element_path=os.path.join(project_dir, element_dir, os.path.basename(element))
                    myfile.write(str(final_element_path)+',')
                myfile.write('\t')
            myfile.write('\n')
            for item in items_list:
                myfile.write(str(item.photo_choice.GetSelection())+'\t')
            myfile.write('\n')
            for item in items_list:
                myfile.write(str(item.check_box.GetValue())+'\t')
            myfile.write('\n')
            for item in items_list:
                myfile.write(str(item.no_foto_cb.GetValue())+'\t')
                
        shutil.copy(path,protocol_path)
        
        for item in items_list:
            if item.old_photos:
                for old_photo in item.old_photos:
                    saved_photos.append(old_photo)
            
    def save_as(self,event):
        global protocol_name, protocol_path, small_photos_dir_path, used_photos_dir, protocol_saved, saved_photos, photo_directory
        protocol_saved=True
        save_as_path=os.path.join(save_path, protocol_name)
        photo_per_page_item=self.photo_per_page.GetMenuItems()
        dlg = wx.FileDialog(
            self, message="Zapisz jako ...", 
            defaultDir=save_as_path, 
            defaultFile=protocol_name, wildcard='Pliki btp (*.btp)|*.btp', style=wx.FD_SAVE|wx.FD_OVERWRITE_PROMPT)
        if dlg.ShowModal() == wx.ID_OK:
            path = dlg.GetPath()
            
            protocol_name=os.path.splitext(os.path.basename(path))[0]     
            protocol_path=os.path.join(work_dir, protocol_name)   
            small_photos_dir_path=os.path.join(protocol_path,'zmniejszone')
            used_photos_dir=os.path.join(protocol_path,'użyte')
            
            main_frame.SetTitle('Lista usterek '+'- '+protocol_name)

            for item_nr in photo_per_page_item:
                nr_bool=self.photo_per_page.IsChecked(item_nr.GetId())
                if nr_bool == True:
                    photo_radio_id=item_nr.GetId()

            with open(path, "w") as myfile:
                myfile.write(photo_directory+'\n')
                myfile.write(str(photo_radio_id)+'\n')
                myfile.write(str(self.description_yes.IsChecked())+'\n\n')
            
                for item in items_list:
                    myfile.write(item.text_panel.GetLineText(0)+'\t')
                myfile.write('\n')
                for item in items_list:
                    for element in item.old_photos:
#                         element_dir=os.path.basename(os.path.dirname(element))
#                         project_dir=os.path.basename(os.path.dirname(os.path.dirname(element)))
#                         final_element_path=os.path.join(project_dir, element_dir, os.path.basename(element))
                    
                        new_element=os.path.join(protocol_name, 'użyte', os.path.basename(element))
#                         new_element=os.path.join(used_photos_dir, os.path.basename(element))
                        myfile.write(str(new_element)+',')
                    myfile.write('\t')
                myfile.write('\n')
                for item in items_list:
                    myfile.write(str(item.photo_choice.GetSelection())+'\t')
                myfile.write('\n')
                for item in items_list:
                    myfile.write(str(item.check_box.GetValue())+'\t')
                myfile.write('\n')
                for item in items_list:
                    myfile.write(str(item.no_foto_cb.GetValue())+'\t')
                    
            if not os.path.exists(protocol_path):
                os.makedirs(protocol_path)
            if not os.path.exists(small_photos_dir_path):
                os.makedirs(small_photos_dir_path)
            if not os.path.exists(used_photos_dir):
                os.makedirs(used_photos_dir)
            shutil.copy(path,protocol_path)
            for item in items_list:
                if item.old_photos:
                    saved_photos.append(item.old_photos)
                    for element in item.old_photos:
                        try: shutil.copy(element, used_photos_dir)
                        except: pass
        dlg.Destroy()
    
    def open_old(self, event):
        global save_path, protocol_name, items_list, protocol_path, small_photos_dir_path, used_photos_dir, protocol_saved, saved_photos, photo_directory
        global work_dir
        
        if protocol_saved==False and protocol_name is not'':
            open_old_message='Poprzednia lista niezapisana. Czy zapisać przed rozpoczęciem nowej?'
            open_old_dlg=wx.MessageBox(open_old_message, 'Informacja', wx.YES_NO|wx.CANCEL) 
            if open_old_dlg==wx.YES:
                self.save_all(event)
            elif open_old_dlg==wx.CANCEL:
                return
        
        protocol_saved=True

        dlg = wx.FileDialog(
            self, message="Wybierz plik",
            defaultDir=save_path, 
            defaultFile="",
            wildcard='Pliki btp (*.btp)|*.btp',
            style=wx.FD_OPEN| wx.FD_CHANGE_DIR|wx.FD_FILE_MUST_EXIST
            )
        if dlg.ShowModal() == wx.ID_OK:
            
            items_list=[]
            children_list=self.scroll_sizer.GetChildren()
            aa=len(children_list)
            for i in range(0,aa):
                self.scroll_sizer.Hide(0)
                self.scroll_sizer.Remove(0)
                self.scroll_panel.Layout()
                
            enable_save()
            enable_flaw()
            enable_start()
            path = dlg.GetPaths()[0]
            protocol_name=os.path.splitext(os.path.basename(path))[0]
            main_frame.SetTitle('Lista usterek '+'- '+protocol_name)
            protocol_path=os.path.join(work_dir, protocol_name)
            small_photos_dir_path=os.path.join(protocol_path,'zmniejszone')
            used_photos_dir=os.path.join(protocol_path,'użyte')
            saved_photos_names=os.listdir(used_photos_dir)
            for photo_name in saved_photos_names:
                photo_name=os.path.join(used_photos_dir, photo_name)
                saved_photos.append(photo_name)
        
#         save_path=os.path.join(final_dir,'file.txt')
            p_paths=[]
            with open(path) as old_file:
                data = old_file.read()
                data = data.split('\n')
                photo_directory=data[0]
                photo_radio_id=data[1]
                yes_no=data[2]
                text=data[4].split('\t')[:-1]
                phts=data[5].split('\t')[:-1]
                grade=data[6].split('\t')[:-1]
                bold=data[7].split('\t')[:-1]
                no_photo=data[8].split('\t')[:-1]
                
                if not os.path.exists(photo_directory):
                    photo_directory=os.path.join(protocol_path, 'zdjęcia do protokołu')
                    os.makedirs(photo_directory)

                for i in range (len(phts)):
                    pths_p_temp=[]
                    try:
                        pths_p=phts[i].split(',')[:-1]
                    except:
                        pass
                     
                    for single_photo in pths_p:
                        pths_p_temp.append(os.path.join(work_dir, single_photo))
                        
                    phts[i]=pths_p_temp

                    new=single(self.scroll_panel)
                    new.text_panel.SetValue(text[i])
                    new.old_photos=phts[i]
                    new.photo_choice.SetSelection(int(grade[i]))
                    if bold[i]=='True':new.check_box.SetValue(True)
                    else:new.check_box.SetValue(False)
                    if no_photo[i]=='True':
                        new.no_foto_cb.SetValue(True)
                        new.photo_button.Disable()
                        new.check_button.Disable()
                    else:new.no_foto_cb.SetValue(False)
                    self.scroll_sizer.Add(new, flag=wx.DOWN, border=5)
                    items_list.append(new)
                    self.main_panel.Layout()
            self.photo_per_page.Check(int(photo_radio_id), True)
            if yes_no=='True':
                self.description_yes.Check(True)
                self.optionsMenu.SetLabel(ID_D, 'Podpis pod zdjęciem: tak')
            else:
                self.description_no.Check(True)
                self.optionsMenu.SetLabel(ID_D, 'Podpis pod zdjęciem: nie')
                
        dlg.Destroy()

                
    def start_project(self, event):
        if protocol_saved==False and protocol_name is not'':
            open_old_message='Poprzednia lista niezapisana. Czy zapisać przed rozpoczęciem nowej?'
            open_old_dlg=wx.MessageBox(open_old_message, 'Informacja', wx.YES_NO|wx.CANCEL) 
            if open_old_dlg==wx.YES:
                self.save_all(event)
            elif open_old_dlg==wx.CANCEL:
                return
        new=new_class(main_frame)
        
                           
class progress_window(wx.Dialog):
    def __init__(self, parent):
        super(progress_window, self).__init__(parent, id=wx.NewId(), title="Postęp",size=(220,260))
        
        self.UI()
        self.Center()
        self.Show()
        
    def UI(self):
        dlg_sizer=wx.BoxSizer(wx.VERTICAL)
        dlg_panel=wx.Panel(self, size=(220,220))
        dlg_panel_sizer=wx.BoxSizer(wx.VERTICAL)
        
        dlg_list_text=wx.StaticText(dlg_panel, label='Wykonywanie listy usterek', style=wx.ALIGN_CENTER, size=(200,20))
        self.list_progress = wx.Gauge(dlg_panel, range=1, size=(200,20), style=wx.GA_HORIZONTAL|wx.GA_SMOOTH)
        self.list_progress_text=wx.StaticText(dlg_panel, label='0/0', style=wx.ALIGN_CENTER, size=(70,30))
        dlg_photo_text=wx.StaticText(dlg_panel, label='Wykonywanie tabeli ze zdjęciami', style=wx.ALIGN_CENTER, size=(200,20))
        self.photo_progress = wx.Gauge(dlg_panel, range=1, size=(200,20), style=wx.GA_HORIZONTAL|wx.GA_SMOOTH)
        self.photo_progress_text=wx.StaticText(dlg_panel, label='0/0', style=wx.ALIGN_CENTER, size=(70,30))
        
        self.dlg_info_text=wx.StaticText(dlg_panel, label='', style=wx.ALIGN_CENTER, size=(200,40))
        
        
        ok_button = wx.Button(dlg_panel, label='OK', size=(50, 30))
        ok_button.Bind(wx.EVT_BUTTON, self.ok_clicked)
        
        dlg_panel_sizer.Add(dlg_list_text)
        dlg_panel_sizer.Add(self.list_progress)
        dlg_panel_sizer.Add(self.list_progress_text, flag=wx.ALIGN_CENTER)
        dlg_panel_sizer.Add(dlg_photo_text)
        dlg_panel_sizer.Add(self.photo_progress)
        dlg_panel_sizer.Add(self.photo_progress_text, flag=wx.ALIGN_CENTER)
        dlg_panel_sizer.Add(self.dlg_info_text, flag=wx.ALIGN_CENTER)
        dlg_panel_sizer.Add(ok_button, flag=wx.ALIGN_CENTER)
        
        dlg_panel.SetSizer(dlg_panel_sizer)
        dlg_sizer.Add(dlg_panel, flag=wx.ALL, border=10)
        self.SetSizer(dlg_sizer)
    
        
    
    def ok_clicked(self, event):
        self.Close()

class new_class(wx.Dialog):
    def __init__(self, parent):
        super(new_class, self).__init__(parent, id=wx.NewId(), title="Nowa lista",size=(220,150))
        
        self.UI()
        self.Center()
        self.Show()
        
    def UI(self):
        dlg_sizer=wx.BoxSizer(wx.VERTICAL)
        dlg_panel=wx.Panel(self, size=(280,180))
        dlg_panel_sizer=wx.BoxSizer(wx.VERTICAL)
        button_sizer=wx.BoxSizer(wx.HORIZONTAL)
        
        dlg_list_text=wx.StaticText(dlg_panel, label='Podaj nazwę folderu roboczego', style=wx.ALIGN_CENTER, size=(200,20))
        self.new_name=wx.TextCtrl(dlg_panel, size=(200,30))
        
        ok_button = wx.Button(dlg_panel, label='OK', size=(50, 30))
        cancel_button = wx.Button(dlg_panel, label='Anluluj', size=(50, 30))
        ok_button.Bind(wx.EVT_BUTTON, self.ok_clicked)
        cancel_button.Bind(wx.EVT_BUTTON, self.cancel_button_clicked)
#         ok_button.SetFocus()
        button_sizer.Add(ok_button, flag=wx.ALIGN_CENTER|wx.TOP, border=10)
        button_sizer.Add(cancel_button, flag=wx.ALIGN_CENTER|wx.TOP, border=10)
        
        dlg_panel_sizer.Add(dlg_list_text)
        dlg_panel_sizer.Add(self.new_name, flag=wx.ALIGN_CENTER)
        dlg_panel_sizer.Add(button_sizer, flag=wx.ALIGN_CENTER)
        
        dlg_panel.SetSizer(dlg_panel_sizer)
        dlg_sizer.Add(dlg_panel, flag=wx.ALL, border=10)
        self.SetSizer(dlg_sizer)
    
        
    def cancel_button_clicked(self, event):
        self.Close()
        
    def ok_clicked(self, event):
        global small_photos_dir_path, used_photos_dir, protocol_name, protocol_path, protocol_saved, items_list
        items_list=[]

        protocol_name=self.new_name.GetLineText(0)
        if protocol_name=='':
            warning_message='Nie podano nazwy.'
            dlg=wx.MessageDialog(self, warning_message, 'Informacja', wx.OK| wx.ICON_INFORMATION)
            dlg.ShowModal()
            pass
        elif protocol_name[0]==' ':
            warning_message='Nazwa nie może rozpoczynać sie od spacji.'
            dlg=wx.MessageDialog(self, warning_message, 'Informacja', wx.OK| wx.ICON_INFORMATION)
            dlg.ShowModal()
            pass
        else:
            protocol_path=os.path.join(work_dir, protocol_name)
            small_photos_dir_path=os.path.join(protocol_path,'zmniejszone')
            used_photos_dir=os.path.join(protocol_path,'użyte')
    
            if os.path.exists(protocol_path):
                warning_message='Folder listy o takiej nazwie istnieje. Czy chcesz nadpisać?'
                dlg=wx.MessageDialog(self, warning_message, 'Ostrzeżenie', wx.YES_NO | wx.CANCEL | wx.ICON_WARNING)
                answ=dlg.ShowModal()
    #             dlg.SetYesNoCancelLabels('Tak', 'Nie','Zmień nazwę')
                if answ==wx.ID_NO:
                    protocol_name=''
                    protocol_path=''
                    small_photos_dir_path=''
                    small_photos_dir_path=''
                    self.Destroy()
                    return
                elif answ==wx.ID_CANCEL:
                    protocol_name=''
                    protocol_path=''
                    small_photos_dir_path=''
                    small_photos_dir_path=''
    #                 self.Destroy()
                    return 
            elif not os.path.exists(protocol_path):
                os.makedirs(protocol_path)
            if not os.path.exists(small_photos_dir_path):
                os.makedirs(small_photos_dir_path)
            if not os.path.exists(used_photos_dir):
                os.makedirs(used_photos_dir)
            protocol_saved=False   
            enable_flaw()
            main_frame.SetTitle('Lista usterek '+'- '+protocol_name)
            
            items_list=[]
            children_list=main_frame.scroll_sizer.GetChildren()
            aa=len(children_list)
            for i in range(0,aa):
                main_frame.scroll_sizer.Hide(0)
                main_frame.scroll_sizer.Remove(0)
                main_frame.scroll_panel.Layout()
            self.Close()            
            
         
class single(wx.Panel):
    def __init__(self, parent):
        super(single, self).__init__(parent)
        
        global items_list, protocol_saved
        self.old_photos=[]
        self.ID=wx.NewId()
        self.parent=parent
        self.main_panel()
        
    def main_panel(self):
        self.single_panel=wx.Panel(self, id=self.ID, size=(size_x-70,40), style=wx.SIMPLE_BORDER)
        self.single_sizer=wx.FlexGridSizer(1,9,10,10)
        
        self.nr_text=wx.StaticText(self.single_panel, label='', style=wx.ST_NO_AUTORESIZE|wx.ALIGN_CENTER, size=(20,20))
        self.text_panel=wx.TextCtrl(self.single_panel, size=flaw_cell_size)
        self.photo_button = wx.Button(self.single_panel, label='Dodaj\nzdjęcia', size=(70,40))
        grade_list=['1','2','3', '4']
        self.photo_choice = wx.Choice(self.single_panel, choices=grade_list, size=(70, 40))
        self.up_button = wx.Button(self.single_panel, label='Góra', size=(70,20))
        self.down_button = wx.Button(self.single_panel, label='Dół', size=(70,20))
        
        self.check_box_panel=wx.Panel(self.single_panel, size=(70,20))
        self.check_box_sizer=wx.BoxSizer(wx.VERTICAL)
        self.check_box=wx.CheckBox(self.check_box_panel, id=wx.NewId(), size=(20,20))
        self.check_box_sizer.Add(self.check_box, flag=wx.ALIGN_CENTER)
        self.check_box_panel.SetSizer(self.check_box_sizer)
        self.check_box_panel.Layout()
        
        self.no_foto_panel=wx.Panel(self.single_panel, size=(70,20))
        self.no_foto_panel_sizer=wx.BoxSizer(wx.VERTICAL)
        self.no_foto_cb=wx.CheckBox(self.no_foto_panel, id=wx.NewId(), size=(20,20))
        self.no_foto_panel_sizer.Add(self.no_foto_cb, flag=wx.ALIGN_CENTER)
        self.no_foto_panel.SetSizer(self.no_foto_panel_sizer)
        self.no_foto_panel.Layout()
        
        self.check_box.Bind(wx.EVT_CHECKBOX, self.check_choice_state)
        self.photo_choice.Bind(wx.EVT_CHOICE, self.check_choice_state)
        self.no_foto_cb.Bind(wx.EVT_CHECKBOX, self.no_photo_dis_en)
        
        self.check_button = wx.Button(self.single_panel, label='Sprawdź\nzdjęcia', size=(70,40))
        
#         self.block_button = wx.Button(self.single_panel, label='Zablokuj', size=(70,20))

        
#         self.no_foto_cb=wx.CheckBox(self.no_foto_panel, id=wx.NewId(), size=(20,20))
        
        
        self.photo_button.Bind(wx.EVT_BUTTON, self.choose_photos)
        self.up_button.Bind(wx.EVT_BUTTON, self.move_up)
        self.down_button.Bind(wx.EVT_BUTTON, self.move_down)
#         self.block_button.Bind(wx.EVT_BUTTON, self.disable_enable)
        self.check_button.Bind(wx.EVT_BUTTON, self.check_photos)
        
        self.single_sizer.Add(self.nr_text, flag=wx.ALIGN_CENTER_VERTICAL)
        self.single_sizer.Add(self.text_panel)
        self.single_sizer.Add(self.photo_button, flag=wx.ALIGN_CENTER_VERTICAL)
        self.single_sizer.Add(self.photo_choice)
        self.single_sizer.Add(self.up_button, flag=wx.ALIGN_CENTER_VERTICAL)
        self.single_sizer.Add(self.down_button, flag=wx.ALIGN_CENTER_VERTICAL)
        self.single_sizer.Add(self.check_box_panel, flag=wx.ALIGN_CENTER_VERTICAL)
        self.single_sizer.Add(self.check_button, flag=wx.ALIGN_CENTER_VERTICAL)
        self.single_sizer.Add(self.no_foto_panel, flag=wx.ALIGN_CENTER_VERTICAL)
        
        self.nr_text.SetLabel(str(len(items_list)+1))
        
        self.SetSizer(self.single_sizer)
        self.photos_paths=[]
    
    def no_photo_dis_en(self, event):
        global protocol_saved
        protocol_saved=False
        if self.no_foto_cb.GetValue():
            self.photo_button.Disable()
            self.check_button.Disable()
            if self.old_photos:
                for old_photo in self.old_photos:
                    shutil.move(old_photo,photo_directory)
                self.old_photos=[]
        elif not self.no_foto_cb.GetValue():
            self.photo_button.Enable()
            self.check_button.Enable()
       
    def check_choice_state(self, event):
        global protocol_saved
        protocol_saved=False
        
    def disable_enable(self, event):
        if self.block_button.GetLabel()=='Zablokuj':
            self.block_button.SetLabel('Odblokuj')
            self.text_panel.Disable()
            self.photo_button.Disable()
            self.photo_choice.Disable()
            self.up_button.Disable()
            self.down_button.Disable()
            self.check_box.Disable()
        elif self.block_button.GetLabel()=='Odblokuj':
            self.block_button.SetLabel('Zablokuj')
            self.text_panel.Enable()
            self.photo_button.Enable()
            self.photo_choice.Enable()
            self.up_button.Enable()
            self.down_button.Enable()
            self.check_box.Enable()
        main_frame.confirmation_text.SetLabel('')
        
    def choose_photos(self, event):
        global photo_directory, protocol_saved
        protocol_saved=False
        self.photos_paths=[]
        if self.old_photos:
                for old_photo in self.old_photos:
                    shutil.move(old_photo,photo_directory)
                self.old_photos=[]
        dlg = wx.FileDialog(
            self, message="Wybierz zdjęcie",
            defaultDir=photo_directory, 
            defaultFile="",
            wildcard='Pliki JPEG (*.jpg)|*.jpg',
            style=wx.FD_OPEN| wx.FD_CHANGE_DIR|wx.FD_FILE_MUST_EXIST|wx.FD_MULTIPLE
            )
        if dlg.ShowModal() == wx.ID_OK:
            self.photos_paths = dlg.GetPaths()
            photo_directory=os.path.dirname(self.photos_paths[0])
            for path in self.photos_paths:
                shutil.move(path,used_photos_dir)
                self.old_photos.append(os.path.join(used_photos_dir,os.path.basename(path)))
            photos_nr=len(self.photos_paths)
            if photos_nr==1:
                conf_text='zdjęcie'
            elif photos_nr>=2 and photos_nr<=4:
                conf_text='zdjęcia'  
            else:
                conf_text='zdjęć'
            main_frame.confirmation_text.SetLabel('Dodano %s\n'%len(self.photos_paths) + conf_text)
            main_frame.confirmation_text.SetBackgroundColour('#c7c7c7')
        dlg.Destroy()       
        
    def move_up(self, event):
        global protocol_saved
        protocol_saved=False
        children_list=main_frame.scroll_sizer.GetChildren()
        aa=len(children_list)
        item_index=items_list.index(self)
        
        if item_index>0:

            temp=items_list[item_index-1]
            items_list[item_index-1]=items_list[item_index]
            items_list[item_index]=temp
            
            for i in range(0,aa):
                main_frame.scroll_sizer.Detach(0)
            
            for j in items_list:
                main_frame.scroll_sizer.Add(j, flag=wx.DOWN, border=5)
                j.nr_text.SetLabel(str(items_list.index(j)+1))

            main_frame.scroll_panel.SetSizer(main_frame.scroll_sizer)
            main_frame.scroll_panel.Layout()
            
            main_frame.confirmation_text.SetLabel('')
            main_frame.confirmation_text.SetBackgroundColour('#f0f0f0')
        else:
            pass

    def move_down(self, event):
        global protocol_saved
        protocol_saved=False
        children_list=main_frame.scroll_sizer.GetChildren()
        aa=len(children_list)
        item_index=items_list.index(self)
        
        if item_index<aa-1:

            temp=items_list[item_index+1]
            items_list[item_index+1]=items_list[item_index]
            items_list[item_index]=temp
            
            for i in range(0,aa):
                main_frame.scroll_sizer.Detach(0)
            
            for j in items_list:
                main_frame.scroll_sizer.Add(j, flag=wx.DOWN, border=5)
                j.nr_text.SetLabel(str(items_list.index(j)+1))

            main_frame.scroll_panel.SetSizer(main_frame.scroll_sizer)
            main_frame.scroll_panel.Layout()
            
            main_frame.confirmation_text.SetLabel('')
            main_frame.confirmation_text.SetBackgroundColour('#f0f0f0')
        else:
            pass
        
    def check_photos(self, event):
        if len(self.old_photos)>0:
            gallery(self, self.text_panel.GetValue()[:20]+'...')
        else:
            no_photo_message='Brak zdjęć w usterce.'
            no_photo_dlg=wx.MessageBox(no_photo_message, 'Informacja', wx.OK|wx.ICON_WARNING) 

galery_size_x=320+16
gallery_size_y=240+95

class gallery(wx.Frame):
    
    def __init__(self, parent, title):
        super(gallery, self).__init__(parent, title=title, size=(galery_size_x, gallery_size_y))
            
        global photo_directory, data_path 
        self.parent=parent
        self.current_picture=0
        
        self.InitUI()    
        self.Centre()
        self.Show() 

        
    def InitUI(self):    
        
        main_panel=wx.Panel(self)
#         main_panel.SetBackgroundColour('#4f5049')
        
        self.photo_panel=wx.Panel(main_panel, size=(320, 240))
#         self.photo_panel.SetBackgroundColour('red')
        buttons_panel=wx.Panel(main_panel)
#         buttons_panel.SetBackgroundColour('blue')
        
        self.main_sizer=wx.BoxSizer(wx.VERTICAL)
        self.photo_count_text=wx.StaticText(main_panel, label='0/0', size=(30,15), style=wx.ALIGN_CENTER)
        self.photo_count_text.SetLabel('1/'+str(len(self.parent.old_photos)))
        self.photo_count_text.SetBackgroundColour('#ededed')
        
        button_panel_sizer=wx.BoxSizer(wx.HORIZONTAL)
        
        os.chdir(data_path)
        left_bmp = wx.Bitmap("left.png", wx.BITMAP_TYPE_ANY)
        delete_bmp = wx.Bitmap("delete.png", wx.BITMAP_TYPE_ANY)
        right_bmp = wx.Bitmap("right.png", wx.BITMAP_TYPE_ANY)
        add_bmp = wx.Bitmap("add.png", wx.BITMAP_TYPE_ANY)
        
        self.button_spacer_left=wx.StaticText(buttons_panel, label='', size=(70,15))
        self.left_button = wx.BitmapButton(buttons_panel, id=BT_LFT_ID, bitmap=left_bmp,size=(left_bmp.GetSize()))
        self.reject_button = wx.BitmapButton(buttons_panel, id=BT_DEL_ID, bitmap=delete_bmp,size=(delete_bmp.GetSize()))
        self.right_button = wx.BitmapButton(buttons_panel, id=BT_RHT_ID, bitmap=right_bmp,size=(right_bmp.GetSize()))
        self.button_spacer_right=wx.StaticText(buttons_panel, label='', size=(40,15))
        self.add_button = wx.BitmapButton(buttons_panel, id=BT_ADD_ID, bitmap=add_bmp,size=(right_bmp.GetSize()))
        
        if len(self.parent.old_photos)<=1:
            self.right_button.Disable()
        self.left_button.Disable()
        
        self.left_button.Bind(wx.EVT_BUTTON, self.prev_photo)
        self.reject_button.Bind(wx.EVT_BUTTON, self.delete_photo)
        self.right_button.Bind(wx.EVT_BUTTON, self.next_photo)
        self.add_button.Bind(wx.EVT_BUTTON, self.add_photo)
        
        
        button_panel_sizer.Add(self.button_spacer_left)
        button_panel_sizer.Add(self.left_button)
        button_panel_sizer.Add(self.reject_button, flag=wx.LEFT|wx.RIGHT, border=10)
        button_panel_sizer.Add(self.right_button)
        button_panel_sizer.Add(self.button_spacer_right)
        button_panel_sizer.Add(self.add_button)
        
        self.main_sizer.Add(self.photo_panel)
        self.main_sizer.Add(self.photo_count_text, flag=wx.ALIGN_CENTER|wx.TOP, border=5)
        self.main_sizer.Add(buttons_panel, flag=wx.ALIGN_CENTER|wx.TOP, border=5)
        
        buttons_panel.SetSizer(button_panel_sizer)
        main_panel.SetSizer(self.main_sizer)

        img = wx.Image(320,240)
        self.imageCtrl = wx.StaticBitmap(self.photo_panel, wx.ID_ANY, wx.Bitmap(img))
        
        self.PhotoMaxSize = 320
        img = wx.Image(self.parent.old_photos[self.current_picture], wx.BITMAP_TYPE_ANY)
        W = img.GetWidth()
        H = img.GetHeight()
        if W > H:
            NewW = self.PhotoMaxSize
            NewH = self.PhotoMaxSize * H / W
        else:
            NewH = self.PhotoMaxSize
            NewW = self.PhotoMaxSize * W / H
        img = img.Scale(NewW,NewH)
        
        
        self.imageCtrl.SetBitmap(wx.Bitmap(img))
        
    def load_image(self, position_in_list):
        img = wx.Image(self.parent.old_photos[position_in_list], wx.BITMAP_TYPE_ANY)
        W = img.GetWidth()
        H = img.GetHeight()
        if W > H:
            NewW = self.PhotoMaxSize
            NewH = self.PhotoMaxSize * H / W
        else:
            NewH = self.PhotoMaxSize
            NewW = self.PhotoMaxSize * W / H
        img = img.Scale(NewW,NewH)
        
        self.imageCtrl.SetBitmap(wx.Bitmap(img))
        self.photo_panel.Refresh()

    def next_photo(self, event):
        myobject = event.GetEventObject()
        self.current_picture+=1
        self.photo_count_text.SetLabel(str(self.current_picture+1)+'/'+str(len(self.parent.old_photos)))
        if self.current_picture==len(self.parent.old_photos)-1:
            myobject.Disable()
        self.load_image(self.current_picture)
        self.left_button.Enable()
        
    def prev_photo(self, event):
        myobject = event.GetEventObject()
        self.current_picture-=1
        self.photo_count_text.SetLabel(str(self.current_picture+1)+'/'+str(len(self.parent.old_photos)))
        if self.current_picture==0:
            myobject.Disable()
        self.load_image(self.current_picture)
        self.right_button.Enable()
    
    def delete_photo(self, event):
        global protocol_saved
        protocol_saved=False
        
        shutil.move(self.parent.old_photos[self.current_picture],photo_directory)
        self.parent.old_photos.pop(self.current_picture)
        
        self.current_picture-=1
        if self.current_picture==-1: self.current_picture=0 
        self.photo_count_text.SetLabel(str(self.current_picture+1)+'/'+str(len(self.parent.old_photos)))
        if self.current_picture==0:
            self.left_button.Disable()
        if len(self.parent.old_photos)==1:
            self.right_button.Disable()
        if len(self.parent.old_photos)>0:
            self.load_image(self.current_picture)
        else:
            self.photo_count_text.SetLabel('0/0')
            self.reject_button.Disable()
            img = wx.Image(320,240)
            os.chdir(data_path)
            img = wx.Image('no_photos.png', wx.BITMAP_TYPE_ANY)
            self.imageCtrl.SetBitmap(wx.Bitmap(img))
            self.photo_panel.Refresh()

        
    def add_photo(self, event):
        global photo_directory, protocol_saved, used_photos_dir
        protocol_saved=False

        dlg = wx.FileDialog(
            self, message="Wybierz zdjęcie",
            defaultDir=photo_directory, 
            defaultFile="",
            wildcard='Pliki JPEG (*.jpg)|*.jpg',
            style=wx.FD_OPEN| wx.FD_CHANGE_DIR|wx.FD_FILE_MUST_EXIST|wx.FD_MULTIPLE
            )
        if dlg.ShowModal() == wx.ID_OK:
            self.photos_paths = dlg.GetPaths()
            photo_directory=os.path.dirname(self.photos_paths[0])
            
            
            for path in self.photos_paths:
                shutil.move(path,used_photos_dir)
                self.parent.old_photos.append(os.path.join(used_photos_dir,os.path.basename(path)))
         
            self.current_picture=len(self.parent.old_photos)-1      
            self.photo_count_text.SetLabel(str(self.current_picture+1)+'/'+str(len(self.parent.old_photos)))
            self.load_image(-1)
            self.reject_button.Enable()
        
        if len(self.parent.old_photos)>1:
            self.left_button.Enable()

        dlg.Destroy()    

def flaws_list():
    global list_1st_photo, list_2nd_photo, list_3rd_photo, list_4th_photo, photos, photo_count, list_progress_count, prgrs_wndw, result_dir
    global dir_path, protocol_path, protocol_name, flaw_list
    global list_1st_txt, list_2nd_txt, list_3rd_txt, list_4th_txt, data_path
    global list_1st_no_photo_txt, list_2nd_no_photo_txt, list_3rd_no_photo_txt, list_4th_no_photo_txt
    document = Document(os.path.join(data_path, 'pattern.docx'))
    sections = document.sections
    section = sections[-1]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = section.footer_distance = Cm(0)
    
    level_one_flaws_index = 0
    p1 = document.add_paragraph()
    p1.style = document.styles['Body Text']
    p1.paragraph_format.space_before=Pt(0)
    p1.paragraph_format.space_after=Pt(0)
    p1.paragraph_format.left_indent = Inches(0.49)
    p1.style.font.size = Pt(10)
    p1.style.font.name = 'Arial'

    p1.add_run('I stopień pilności:\n').bold = True
    
    flaw_len=len(list_1st_txt)+len(list_2nd_txt)+len(list_3rd_txt)+len(list_4th_txt)+len(list_1st_no_photo_txt)+len(list_2nd_no_photo_txt)+len(list_3rd_no_photo_txt)+len(list_4th_no_photo_txt)
    
    if len(list_1st_txt)==0 and len(list_1st_no_photo_txt)==0:
        p1.add_run('\nBrak zaleceń remontowych.')
    
    for flaw in list_1st_txt:
        list_progress_count+=1
        prgrs_wndw.list_progress_text.SetLabel(str(list_progress_count)+'/'+str(flaw_len))
        prgrs_wndw.list_progress.SetValue(list_progress_count)
        
        p_l1 = document.add_paragraph()
        p_l1.style = document.styles['Body Text']
        p_l1.paragraph_format.space_before=Pt(0)
        p_l1.paragraph_format.space_after=Pt(0)
        p_l1.paragraph_format.left_indent = Inches(0.74)
        p_l1.style.font.size = Pt(10)
        p_l1.style.font.name = 'Arial'
        
        start_point=len(photos)
        for item in items_list:
            if item.text_panel.GetLineText(0)==flaw:
                photos+=item.old_photos
                p_l1_run=p_l1.add_run('%s. ' %str(level_one_flaws_index+1) + flaw)
                p_l1_run.bold = item.check_box.GetValue()
                p_l1_run.underline = item.check_box.GetValue()
                level_one_flaws_index+=1
        p_l1.add_run(' (zdj. nr ')
        
        for i in range(start_point+1, len(photos)):
            if len(photos)-start_point==2: #only if list is two elements long
                p_l1.add_run('%s,' %i)
            elif i-start_point==2: #if a second element appears in list
                p_l1.add_run('%s' %str(start_point+1))
            flaw_list.append(flaw)
        if len(photos)-start_point>2:
            p_l1.add_run('-')
        p_l1.add_run('%s' %str(len(photos)))
        flaw_list.append(flaw)
        p_l1.add_run(')')
        
    for flaw_np in list_1st_no_photo_txt:
        list_progress_count+=1
        prgrs_wndw.list_progress_text.SetLabel(str(list_progress_count)+'/'+str(flaw_len))
        prgrs_wndw.list_progress.SetValue(list_progress_count)
        
        p_l1_np = document.add_paragraph()
        p_l1_np.style = document.styles['Body Text']
        p_l1_np.paragraph_format.space_before=Pt(0)
        p_l1_np.paragraph_format.space_after=Pt(0)
        p_l1_np.paragraph_format.left_indent = Inches(0.74)
        p_l1_np.style.font.size = Pt(10)
        p_l1_np.style.font.name = 'Arial'
        
        for item in items_list:
            if item.text_panel.GetLineText(0)==flaw_np:
#                 photos+=item.old_photos
                p_l1_run_np=p_l1_np.add_run('%s. ' %str(level_one_flaws_index+1)+ flaw_np)
                p_l1_run_np.bold = item.check_box.GetValue()
                p_l1_run_np.underline = item.check_box.GetValue()
                level_one_flaws_index+=1
    
    level_two_flaws_index = 0
    p2 = document.add_paragraph()
    p2.style = document.styles['Body Text']
    p2.paragraph_format.space_before=Pt(0)
    p2.paragraph_format.space_after=Pt(0)
    p2.paragraph_format.left_indent = Inches(0.49)
    p2.style.font.size = Pt(10)
    p2.style.font.name = 'Arial'
    
    p2.add_run('\nII stopień pilności:\n').bold = True
    
    if len(list_2nd_txt)==0 and len(list_2nd_no_photo_txt)==0:
        p2.add_run('\nBrak zaleceń remontowych.')
    
    for flaw in list_2nd_txt:
        list_progress_count+=1
        prgrs_wndw.list_progress.SetValue(list_progress_count)
        prgrs_wndw.list_progress_text.SetLabel(str(list_progress_count)+'/'+str(flaw_len))
        
        p_l2 = document.add_paragraph()
        p_l2.style = document.styles['Body Text']
        p_l2.paragraph_format.space_before=Pt(0)
        p_l2.paragraph_format.space_after=Pt(0)
        p_l2.paragraph_format.left_indent = Inches(0.74)
        p_l2.style.font.size = Pt(10)
        p_l2.style.font.name = 'Arial'
        
        start_point=len(photos)
        for item in items_list:
            if item.text_panel.GetLineText(0)==flaw:
                photos+=item.old_photos
                p_l2_run=p_l2.add_run('%s. ' %str(level_two_flaws_index+1) + flaw)
                p_l2_run.bold = item.check_box.GetValue()
                p_l2_run.underline = item.check_box.GetValue()
                level_two_flaws_index+=1
        p_l2.add_run(' (zdj. nr ')
        
        for i in range(start_point+1, len(photos)):
            if len(photos)-start_point==2:
                p_l2.add_run('%s,' %i)
            elif i-start_point==2:
                p_l2.add_run('%s' %str(start_point+1))
            flaw_list.append(flaw)
        if len(photos)-start_point>2:
            p_l2.add_run('-')
        flaw_list.append(flaw)
        p_l2.add_run('%s' %str(len(photos)))
        p_l2.add_run(')')
        
    for flaw_np in list_2nd_no_photo_txt:
        list_progress_count+=1
        prgrs_wndw.list_progress_text.SetLabel(str(list_progress_count)+'/'+str(flaw_len))
        prgrs_wndw.list_progress.SetValue(list_progress_count)
        
        p_l2_np = document.add_paragraph()
        p_l2_np.style = document.styles['Body Text']
        p_l2_np.paragraph_format.space_before=Pt(0)
        p_l2_np.paragraph_format.space_after=Pt(0)
        p_l2_np.paragraph_format.left_indent = Inches(0.74)
        p_l2_np.style.font.size = Pt(10)
        p_l2_np.style.font.name = 'Arial'
        
        for item in items_list:
            if item.text_panel.GetLineText(0)==flaw_np:
#                 photos+=item.old_photos
                p_l2_run_np=p_l2_np.add_run('%s. ' %str(level_two_flaws_index+1) + flaw_np)
                p_l2_run_np.bold = item.check_box.GetValue()
                p_l2_run_np.underline = item.check_box.GetValue()
                level_two_flaws_index+=1

    level_three_flaws_index=0
    p3 = document.add_paragraph()
    p3.style = document.styles['Body Text']
    p3.paragraph_format.space_before=Pt(0)
    p3.paragraph_format.space_after=Pt(0)
    p3.paragraph_format.left_indent = Inches(0.49)
    p3.style.font.size = Pt(10)
    p3.style.font.name = 'Arial'
    
    p3.add_run('\nIII stopień pilności:\n').bold = True
    
    if len(list_3rd_txt)==0 and len(list_3rd_no_photo_txt)==0:
        p3.add_run('\nBrak zaleceń remontowych.')
        
    for flaw in list_3rd_txt:
        list_progress_count+=1
        prgrs_wndw.list_progress.SetValue(list_progress_count)
        prgrs_wndw.list_progress_text.SetLabel(str(list_progress_count)+'/'+str(flaw_len))
        
        p_l3 = document.add_paragraph()
        p_l3.style = document.styles['Body Text']
        p_l3.paragraph_format.space_before=Pt(0)
        p_l3.paragraph_format.space_after=Pt(0)
        p_l3.paragraph_format.left_indent = Inches(0.74)
        p_l3.style.font.size = Pt(10)
        p_l3.style.font.name = 'Arial'
        
        start_point=len(photos)
        for item in items_list:
            if item.text_panel.GetLineText(0)==flaw:
                photos+=item.old_photos
                p_l3_run=p_l3.add_run('%s. ' %str(level_three_flaws_index+1) + flaw)
                p_l3_run.bold = item.check_box.GetValue()
                p_l3_run.underline = item.check_box.GetValue()
                level_three_flaws_index+=1
        p_l3.add_run(' (zdj. nr ')
        
        for i in range(start_point+1, len(photos)):
            if len(photos)-start_point==2:
                p_l3.add_run('%s,' %i)
            elif i-start_point==2:
                p_l3.add_run('%s' %str(start_point+1))
            flaw_list.append(flaw)
        if len(photos)-start_point>2:
            p_l3.add_run('-')
        flaw_list.append(flaw)
        p_l3.add_run('%s' %str(len(photos)))
        p_l3.add_run(')')
        
    for flaw_np in list_3rd_no_photo_txt:
        list_progress_count+=1
        prgrs_wndw.list_progress_text.SetLabel(str(list_progress_count)+'/'+str(flaw_len))
        prgrs_wndw.list_progress.SetValue(list_progress_count)
        
        p_l3_np = document.add_paragraph()
        p_l3_np.style = document.styles['Body Text']
        p_l3_np.paragraph_format.space_before=Pt(0)
        p_l3_np.paragraph_format.space_after=Pt(0)
        p_l3_np.paragraph_format.left_indent = Inches(0.74)
        p_l3_np.style.font.size = Pt(10)
        p_l3_np.style.font.name = 'Arial'
        
        for item in items_list:
            if item.text_panel.GetLineText(0)==flaw_np:
#                 photos+=item.old_photos
                p_l3_run_np=p_l3_np.add_run('%s. ' %str(level_three_flaws_index+1) + flaw_np)
                p_l3_run_np.bold = item.check_box.GetValue()
                p_l3_run_np.underline = item.check_box.GetValue()
                level_three_flaws_index+=1

    level_four_flaws_index=0
    p4 = document.add_paragraph()
    p4.style = document.styles['Body Text']
    p4.paragraph_format.space_before=Pt(0)
    p4.paragraph_format.space_after=Pt(0)
    p4.paragraph_format.left_indent = Inches(0.49)
    p4.style.font.size = Pt(10)
    p4.style.font.name = 'Arial'
    
    p4.add_run('\nIV stopień pilności:\n').bold = True
    
    if len(list_4th_txt)==0 and len(list_4th_no_photo_txt)==0:
        p4.add_run('\nBrak zaleceń remontowych.')
        
    for flaw in list_4th_txt:
        list_progress_count+=1
        prgrs_wndw.list_progress.SetValue(list_progress_count)
        prgrs_wndw.list_progress_text.SetLabel(str(list_progress_count)+'/'+str(flaw_len))
        
        p_l4 = document.add_paragraph()
        p_l4.style = document.styles['Body Text']
        p_l4.paragraph_format.space_before=Pt(0)
        p_l4.paragraph_format.space_after=Pt(0)
        p_l4.paragraph_format.left_indent = Inches(0.74)
        p_l4.style.font.size = Pt(10)
        p_l4.style.font.name = 'Arial'
        
        start_point=len(photos)
        for item in items_list:
            if item.text_panel.GetLineText(0)==flaw:
                photos+=item.old_photos
                p_l4_run=p_l4.add_run('%s. ' %str(level_four_flaws_index+1) + flaw)
                p_l4_run.bold = item.check_box.GetValue()
                p_l4_run.underline = item.check_box.GetValue()
                level_four_flaws_index+=1
        p_l4.add_run(' (zdj. nr ')
        
        for i in range(start_point+1, len(photos)):
            if len(photos)-start_point==2:
                p_l4.add_run('%s,' %i)
            elif i-start_point==2:
                p_l4.add_run('%s' %str(start_point+1))
            flaw_list.append(flaw)
        if len(photos)-start_point>2:
            p_l4.add_run('-')
        flaw_list.append(flaw)
        p_l4.add_run('%s' %str(len(photos)))
        p_l4.add_run(')')
        
    for flaw_np in list_4th_no_photo_txt:
        list_progress_count+=1
        prgrs_wndw.list_progress_text.SetLabel(str(list_progress_count)+'/'+str(flaw_len))
        prgrs_wndw.list_progress.SetValue(list_progress_count)
        
        p_l4_np = document.add_paragraph()
        p_l4_np.style = document.styles['Body Text']
        p_l4_np.paragraph_format.space_before=Pt(0)
        p_l4_np.paragraph_format.space_after=Pt(0)
        p_l4_np.paragraph_format.left_indent = Inches(0.74)
        p_l4_np.style.font.size = Pt(10)
        p_l4_np.style.font.name = 'Arial'
        
        for item in items_list:
            if item.text_panel.GetLineText(0)==flaw_np:
#                 photos+=item.old_photos
                p_l4_run_np=p_l4_np.add_run('%s. ' %str(level_four_flaws_index+1) + flaw_np)
                p_l4_run_np.bold = item.check_box.GetValue()
                p_l4_run_np.underline = item.check_box.GetValue()
                level_four_flaws_index+=1

    
    global list_docx_path_copy
    
#     list_docx_path=os.path.join(result_dir, protocol_name+'-lista.docx')
    list_docx_path_copy=os.path.join(protocol_path, protocol_name+'-lista.docx')
#     document.save(list_docx_path)
    document.save(list_docx_path_copy)
    
    
def photo_table():
    global photos, photo_count, protocol_name, small_photos_list, flaw_list, result_dir
    global list_docx_path_copy
    
    document = Document(list_docx_path_copy)
    document.add_page_break()
#     sections = document.sections
#     section = sections[-1]
#     section.orientation = WD_ORIENT.PORTRAIT
#     section.page_width = Cm(21)
#     section.page_height = Cm(29.7)
#     section.left_margin = section.right_margin = Cm(2.5)
#     section.top_margin = Cm(1.5)
#     section.bottom_margin = Cm(2.5)
#     section.header_distance = section.footer_distance = Cm(0)
    
    if main_frame.two_photo.IsChecked():
        new_h=9
        colN=1
        rowN = len(photos)*2
    elif main_frame.three_photo.IsChecked():
        new_h=7
        colN=1
        rowN = len(photos)*2
    elif main_frame.six_photo.IsChecked():
        new_h=6
        colN=2
        rowN = math.ceil(len(photos)/2)*2
    
    table = document.add_table(rows=rowN, cols=colN)
    table.style = document.styles['Table Grid']
    table.autofit = True
    for i in range(0,len(photos)):
        photo_count+=1
        prgrs_wndw.photo_progress.SetValue(photo_count)
        prgrs_wndw.photo_progress_text.SetLabel(str(photo_count)+'/'+str(len(photos)))
        prgrs_wndw.Update()
        if main_frame.six_photo.IsChecked(): image_cell=table.cell(i-i%2,i%2)
        else: image_cell=table.cell(i*2,0)
        
#         print('Zdj %s ' %i, ": %s" %math.floor(i/2), ",%s" %(i%2))
#         print('---')
#         print('Opis %s ' %i, ": %s,%s" %math.floor(i/2) %(i%2))
        
        photo_paragraph = image_cell.paragraphs[0]
        photo_paragraph.style.font.size = Pt(9)
        photo_paragraph.style.font.name = 'Tahoma'
        photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first_run = photo_paragraph.add_run()
        
        photo = Image.open(photos[i])
        w, h=photo.size
        aspect_ratio=w/h
        
        if aspect_ratio==4/3 and w>=640:
            new_width=640
            new_height=480
        else:
            fac=1
            for jj in range(20,101,2):
                if jj/100*w>899 and jj/100*w<1001:
                    fac=jj/100
                    break
                    
            new_width=int(w*fac)
            new_height=int(h*fac)

        photo = photo.resize((new_width, new_height), Image.ANTIALIAS)
        p_name='photo'+str(photo_count)+'.jpg'
        
        p_save_path=os.path.join(small_photos_dir_path,p_name)
#         small_photos_list.append(p_save_path)
        photo=photo.save(p_save_path, 'JPEG', quality=90,optimize=True)
        p_path=os.path.join(os.path.realpath(photo),p_save_path)
        new_w=new_h*aspect_ratio
        first_run.add_picture(p_path, width = Cm(new_w), height = Cm(new_h))
        
        if main_frame.six_photo.IsChecked(): text_cell=table.cell(i-i%2+1,i%2)
        else: text_cell=table.cell(i*2+1,0)
        
        text_paragraph=text_cell.add_paragraph()
        text_paragraph.style.font.name = 'Tahoma'
        text_paragraph.paragraph_format.space_before=Pt(0)
        text_paragraph.paragraph_format.space_after=Pt(6)
        
        
        text_paragraph.style = document.styles['Normal']
        text_paragraph.style.font.size = Pt(9)
        text_paragraph.add_run('Zdjęcie nr %s - ' %(i+1))
        text_paragraph.add_run(flaw_list[i])
    
    photos_docx_path=os.path.join(result_dir, protocol_name+'-wynik.docx')
    photos_docx_path_copy=os.path.join(protocol_path, protocol_name+'-wynik.docx')
    document.save(photos_docx_path)
    document.save(photos_docx_path_copy)
#     os.remove(list_docx_path_copy)
    
def enable_save():
    main_frame.menu_save_button.Enable(True)
    main_frame.menu_save_as_button.Enable(True)
def enable_flaw():
    main_frame.add_next.Enable(True)
    main_frame.remove_last.Enable(True)
def enable_start():
    main_frame.menu_start_button.Enable(True)
    
# def move_back():
#     global photo_directory, dir_path
#     to_move_list=[]
#     files_list=os.listdir(dir_path)
#     
#     if photo_directory is not dir_path:
#         for item in files_list:
#             if item.endswith('.JPG'):
#                 to_move_list.append(os.path.join(dir_path,item))
#         for photo in to_move_list:
#             shutil.move(photo, photo_directory)
    
        
if __name__ == '__main__':
    
    app = wx.App()
    main_frame=main_window(None, title='Lista usterek')
    app.MainLoop()